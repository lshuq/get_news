import time
import requests
from selenium import webdriver
from bs4 import BeautifulSoup
import urllib.request
from docx import Document
from docx.shared import Inches, Pt
import io


# 获取微信公众号新闻

def get_access():
    return webdriver.PhantomJS()


def search_page(driver, url):
    driver.get(url)
    elem = driver.find_element_by_class_name("query")
    elem.send_keys(u"国际农业航空施药技术联合实验室")
    btn = driver.find_element_by_class_name("swz2")
    btn.click()
    time.sleep(2)
    page = driver.page_source
    return page


def load_page(driver, url):
    driver.get(url)
    page = driver.page_source
    return page


def next_page(page):
    soup = BeautifulSoup(page, "html5lib")
    link = soup.find('div', attrs={"class": "img-box"})
    link_to = link.find('a')
    return link_to.attrs['href']


def get_news_urls(page):
    base_url = 'https://mp.weixin.qq.com'
    # 存储新闻页面的列表
    # news_urls = {}
    news_urls = []
    soup = BeautifulSoup(page, "html5lib")
    # 新闻列表
    news_list = soup.find_all('h4', attrs={'class': 'weui_media_title'})
    # 全部新闻标题
    for i in news_list:
        # news_title = i.get_text()
        news_url = base_url + i.attrs['hrefs']
        # news_urls[news_title] = news_url
        news_urls.append(news_url)
    return news_urls


def download_page(url):
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_2) '
                             'AppleWebKit/537.36 (KHTML, like Gecko) Chrome/47.0.2526.80 Safari/537.36'}
    request = urllib.request.Request(url, headers=headers)
    response = urllib.request.urlopen(request)
    data = response.read()
    return data


def get_imgs_url(url):
    images = []
    page = urllib.request.urlopen(url)
    soup = BeautifulSoup(page, "html5lib")
    imgs = soup.find_all('img')
    for img in imgs:
        try:
            images.append(img.attrs['data-src'])
        except:
            pass
    return images


def get_news(news_urls):
    for url in news_urls:
        document = Document()
        html = download_page(url)
        soup = BeautifulSoup(html, 'html.parser')
        content = soup.find('div', attrs={'class': 'rich_media_content'})
        head = soup.find('h2', attrs={'class': 'rich_media_title'})
        document.add_heading(head.get_text(), 0)
        head = soup.find('div', attrs={'class': 'rich_media_meta_list'})
        detail = head.find_all('em')
        txt = ''
        doc_name = head.find('em').get_text()
        for i in detail:
            txt += i.get_text()
        p = document.add_paragraph(txt).add_run()
        p.font.size = Pt(10)
        paras = content.find_all('p')
        for para in paras:
            img = para.find('img')
            if img is not None:
                image = download_page(img.attrs['data-src'])
                img_io = io.BytesIO()
                img_io.write(image)
                img_io.seek(0)
                try:
                    document.add_picture(img_io, width=Inches(4.5))
                except Exception as e:
                    pass
                img_io.close()
            txt = para.get_text()
            document.add_paragraph(txt)
        document.add_page_break()
        document.save('%s.docx' % doc_name)
    print("Done")
    return


def initial():
    html_url = "http://weixin.sogou.com/"
    web_look = get_access()
    html_url = search_page(web_look, html_url)
    true_url = next_page(html_url)
    news_urls = get_news_urls(load_page(web_look, true_url))
    get_news(news_urls)
    web_look.quit()


if __name__ == '__main__':
    initial()
